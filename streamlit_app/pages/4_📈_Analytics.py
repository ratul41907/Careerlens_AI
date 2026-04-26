"""
Analytics Dashboard - Platform Usage Insights
"""
import streamlit as st
import sys
from pathlib import Path
import json
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import random

# Mobile responsiveness - Day 25
import sys
import os

# Add utils path consistently with other pages
utils_path = os.path.join(os.path.dirname(__file__), '..')
if utils_path not in sys.path:
    sys.path.insert(0, utils_path)

try:
    from utils.mobile_styles import inject_mobile_styles
    inject_mobile_styles()
except ImportError:
    pass  # Mobile styles optional
# Add project root to path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

# Page config
st.set_page_config(
    page_title="Analytics - CareerLens AI",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Dark theme CSS
st.markdown("""
<style>
    /* Force dark theme */
    .stApp {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%) !important;
    }
    
    /* Hide branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* All text light colored */
    * {
        color: #e2e8f0 !important;
    }
    
    /* Sidebar */
    [data-testid="stSidebar"] {
        background: rgba(15, 23, 42, 0.95) !important;
    }
    
    /* Metric cards */
    [data-testid="stMetricValue"] {
        font-size: 2rem !important;
        color: #60a5fa !important;
    }
    
    /* Success/Info messages */
    .stSuccess {
        background: rgba(16, 185, 129, 0.1) !important;
        border: 1px solid rgba(16, 185, 129, 0.3) !important;
        color: #10b981 !important;
    }
    
    .stInfo {
        background: rgba(59, 130, 246, 0.1) !important;
        border: 1px solid rgba(59, 130, 246, 0.3) !important;
        color: #3b82f6 !important;
    }
    
    /* Stat box */
    .stat-box {
        background: rgba(30, 41, 59, 0.7);
        border-left: 4px solid #3b82f6;
        border-radius: 8px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    
    .stat-box-success {
        border-left-color: #10b981;
    }
    
    .stat-box-warning {
        border-left-color: #f59e0b;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state for demo data
if 'analytics_data' not in st.session_state:
    # Create match history
    match_history = []
    for i in range(30):
        date = datetime.now() - timedelta(days=29-i)
        score = random.randint(55, 95)
        match_history.append({
            'date': date.strftime('%Y-%m-%d'),
            'score': score
        })
    
    # Create skill gaps
    skill_gaps = {
        'Docker': 15,
        'Kubernetes': 14,
        'AWS': 13,
        'React': 11,
        'MongoDB': 10,
        'Redis': 9,
        'GraphQL': 7,
        'TypeScript': 5
    }
    
    # Generate sample analytics data
    st.session_state.analytics_data = {
        'total_matches': 47,
        'avg_score': 73.5,
        'cvs_generated': 12,
        'interviews_practiced': 23,
        'match_history': match_history,
        'skill_gaps': skill_gaps
    }

# Warning about demo data
st.info("""
💡 **Note:** This dashboard shows demo visualization data. The Home page stats track your actual session activity.
""")

# Title
st.markdown("""
<h1 style="background: linear-gradient(90deg, #60a5fa, #a78bfa); -webkit-background-clip: text; 
           -webkit-text-fill-color: transparent; font-size: 3rem; font-weight: 800; margin-bottom: 0.5rem;">
    📈 Analytics Dashboard
</h1>
<p style="color: #94a3b8 !important; font-size: 1.25rem; margin-bottom: 2rem;">
    Track your progress • Identify skill gaps • Measure success
</p>
""", unsafe_allow_html=True)

# Get analytics data
data = st.session_state.analytics_data

# ============================================================================
# OVERVIEW METRICS
# ============================================================================
st.markdown("### 📊 Overview Metrics")

metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4)

with metric_col1:
    st.metric(
        "Total CV Matches",
        data['total_matches'],
        "+5 this week"
    )

with metric_col2:
    st.metric(
        "Average Match Score",
        f"{data['avg_score']}%",
        "+3.2% vs last month"
    )

with metric_col3:
    st.metric(
        "CVs Generated",
        data['cvs_generated'],
        "+2 this week"
    )

with metric_col4:
    st.metric(
        "Interview Questions",
        data['interviews_practiced'],
        "+8 this week"
    )

st.markdown("---")

# ============================================================================
# MATCH SCORE TRENDS
# ============================================================================
st.markdown("### 📈 Match Score Trends (Last 30 Days)")

# Create DataFrame
match_df = pd.DataFrame(data['match_history'])
match_df['date'] = pd.to_datetime(match_df['date'])

# Create line chart with FIXED COLORS
fig_trend = go.Figure()

fig_trend.add_trace(go.Scatter(
    x=match_df['date'],
    y=match_df['score'],
    mode='lines+markers',
    name='Match Score',
    line=dict(color='#60a5fa', width=3),
    marker=dict(size=8, color='#3b82f6'),
    fill='tozeroy',
    fillcolor='rgba(96, 165, 250, 0.2)'
))

# Add threshold line
fig_trend.add_hline(
    y=75, 
    line_dash="dash", 
    line_color="#10b981",
    annotation_text="Target (75%)",
    annotation_position="right",
    annotation_font_color="#ffffff"
)

fig_trend.update_layout(
    plot_bgcolor='rgba(0,0,0,0)',
    paper_bgcolor='rgba(0,0,0,0)',
    font=dict(color='#ffffff', size=14),
    xaxis=dict(
        showgrid=True,
        gridcolor='rgba(255,255,255,0.1)',
        title='Date',
        titlefont=dict(color='#ffffff', size=16)
    ),
    yaxis=dict(
        showgrid=True,
        gridcolor='rgba(255,255,255,0.1)',
        title='Match Score (%)',
        range=[0, 100],
        titlefont=dict(color='#ffffff', size=16)
    ),
    hovermode='x unified',
    height=400
)

st.plotly_chart(fig_trend, use_container_width=True)

# Insights
trend_col1, trend_col2 = st.columns(2)

with trend_col1:
    st.markdown("""
    <div class="stat-box stat-box-success">
        <h4 style="color: #10b981 !important; margin: 0 0 0.5rem 0;">📈 Improving Trend</h4>
        <p style="color: #ffffff !important; margin: 0;">
            Your match scores have increased by <strong>12%</strong> over the last 30 days. 
            Keep focusing on skill development!
        </p>
    </div>
    """, unsafe_allow_html=True)

with trend_col2:
    st.markdown("""
    <div class="stat-box">
        <h4 style="color: #60a5fa !important; margin: 0 0 0.5rem 0;">🎯 Best Match</h4>
        <p style="color: #ffffff !important; margin: 0;">
            Your highest match score was <strong>92%</strong> for a Senior Backend Engineer 
            position on March 1, 2026.
        </p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# ============================================================================
# SKILL GAP ANALYSIS
# ============================================================================
st.markdown("### 🎯 Top Skill Gaps")

st.info("💡 These are the most frequently missing skills across your CV matches. Focus on learning these to improve your match scores!")

# Create DataFrame for skill gaps
skill_gap_df = pd.DataFrame([
    {'skill': skill, 'frequency': freq}
    for skill, freq in sorted(data['skill_gaps'].items(), key=lambda x: x[1], reverse=True)
])

# Create bar chart with STANDARD COLORS
fig_skills = go.Figure()

fig_skills.add_trace(go.Bar(
    x=skill_gap_df['frequency'],
    y=skill_gap_df['skill'],
    orientation='h',
    marker=dict(
        color=skill_gap_df['frequency'],
        colorscale=[[0, '#3b82f6'], [0.5, '#8b5cf6'], [1, '#ec4899']],
        showscale=False
    ),
    text=skill_gap_df['frequency'],
    textposition='outside',
    textfont=dict(color='#ffffff', size=14)
))

fig_skills.update_layout(
    plot_bgcolor='rgba(0,0,0,0)',
    paper_bgcolor='rgba(0,0,0,0)',
    font=dict(color='#ffffff', size=14),
    xaxis=dict(
        showgrid=True,
        gridcolor='rgba(255,255,255,0.1)',
        title='Times Missing',
        titlefont=dict(color='#ffffff', size=16)
    ),
    yaxis=dict(
        showgrid=False,
        title='',
        tickfont=dict(color='#ffffff', size=14)
    ),
    height=400,
    margin=dict(l=100, r=50, t=50, b=50)
)

st.plotly_chart(fig_skills, use_container_width=True)

# Recommendations
st.markdown("### 💡 Recommended Learning Path")

rec_col1, rec_col2, rec_col3 = st.columns(3)

with rec_col1:
    st.markdown("""
    <div class="stat-box">
        <h4 style="color: #ef4444 !important; margin: 0 0 0.5rem 0;">🔥 High Priority</h4>
        <p style="color: #ffffff !important; margin: 0; font-size: 0.9rem;">
            <strong>Docker</strong> (missing 15x)<br>
            <strong>Kubernetes</strong> (missing 14x)<br>
            <strong>AWS</strong> (missing 13x)
        </p>
    </div>
    """, unsafe_allow_html=True)

with rec_col2:
    st.markdown("""
    <div class="stat-box">
        <h4 style="color: #f59e0b !important; margin: 0 0 0.5rem 0;">⚡ Medium Priority</h4>
        <p style="color: #ffffff !important; margin: 0; font-size: 0.9rem;">
            <strong>React</strong> (missing 11x)<br>
            <strong>MongoDB</strong> (missing 10x)<br>
            <strong>Redis</strong> (missing 9x)
        </p>
    </div>
    """, unsafe_allow_html=True)

with rec_col3:
    st.markdown("""
    <div class="stat-box">
        <h4 style="color: #10b981 !important; margin: 0 0 0.5rem 0;">📚 Nice to Have</h4>
        <p style="color: #ffffff !important; margin: 0; font-size: 0.9rem;">
            <strong>GraphQL</strong> (missing 7x)<br>
            <strong>TypeScript</strong> (missing 5x)<br>
        </p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# ============================================================================
# CV GENERATION STATS
# ============================================================================
st.markdown("### 📝 CV Generation Statistics")

cv_col1, cv_col2 = st.columns([1, 1])

with cv_col1:
    # CV generation modes
    cv_modes = pd.DataFrame({
        'Mode': ['Manual Entry', 'Auto from JD', 'Document Upload'],
        'Count': [7, 3, 2]
    })
    
    fig_cv_modes = go.Figure(data=[go.Pie(
        labels=cv_modes['Mode'],
        values=cv_modes['Count'],
        marker=dict(colors=['#3b82f6', '#8b5cf6', '#ec4899']),
        textfont=dict(color='#ffffff', size=14)
    )])
    
    fig_cv_modes.update_layout(
        title='CVs by Generation Mode',
        title_font=dict(color='#ffffff', size=16),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#ffffff'),
        showlegend=True,
        height=350
    )
    
    st.plotly_chart(fig_cv_modes, use_container_width=True)

with cv_col2:
    # Download formats
    downloads = pd.DataFrame({
        'Format': ['DOCX', 'PDF'],
        'Count': [10, 2]
    })
    
    fig_downloads = go.Figure(data=[go.Pie(
        labels=downloads['Format'],
        values=downloads['Count'],
        marker=dict(colors=['#10b981', '#f59e0b']),
        textfont=dict(color='#ffffff', size=14)
    )])
    
    fig_downloads.update_layout(
        title='Download Format Preference',
        title_font=dict(color='#ffffff', size=16),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#ffffff'),
        showlegend=True,
        height=350
    )
    
    st.plotly_chart(fig_downloads, use_container_width=True)

st.markdown("---")

# ============================================================================
# INTERVIEW PREP STATS
# ============================================================================
st.markdown("### 🎓 Interview Preparation Progress")

interview_col1, interview_col2 = st.columns([2, 1])

with interview_col1:
    # Questions by category - SAME COLORS AS INTERVIEW PREP
    questions_df = pd.DataFrame({
        'Category': ['Behavioral', 'Technical', 'Coding', 'System Design'],
        'Questions': [8, 7, 5, 3]
    })
    
    fig_questions = go.Figure()
    
    fig_questions.add_trace(go.Bar(
        x=questions_df['Category'],
        y=questions_df['Questions'],
        marker=dict(
            color=['#3b82f6', '#8b5cf6', '#ec4899', '#10b981']
        ),
        text=questions_df['Questions'],
        textposition='outside',
        textfont=dict(color='#ffffff', size=14)
    ))
    
    fig_questions.update_layout(
        title='Questions Practiced by Category',
        title_font=dict(color='#ffffff', size=16),
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#ffffff', size=14),
        xaxis=dict(
            showgrid=False,
            tickfont=dict(color='#ffffff', size=14)
        ),
        yaxis=dict(
            showgrid=True,
            gridcolor='rgba(255,255,255,0.1)',
            titlefont=dict(color='#ffffff', size=16)
        ),
        showlegend=False,
        height=350
    )
    
    st.plotly_chart(fig_questions, use_container_width=True)

with interview_col2:
    st.markdown("### 📊 Evaluation Scores")
    
    st.metric("Average Score", "78/100", "+5 points")
    st.metric("Best Score", "92/100", "Excellent")
    st.metric("Answers Evaluated", "8", "+3 this week")
    
    st.markdown("""
    <div class="stat-box stat-box-success" style="margin-top: 1rem;">
        <p style="color: #ffffff !important; margin: 0; font-size: 0.9rem;">
            <strong>🌟 Achievement Unlocked!</strong><br>
            You've evaluated 8+ answers. Keep practicing!
        </p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# ============================================================================
# PLATFORM USAGE SUMMARY
# ============================================================================
st.markdown("### 📅 Platform Usage (Last 7 Days)")

usage_data = pd.DataFrame({
    'Date': [(datetime.now() - timedelta(days=6-i)).strftime('%b %d') for i in range(7)],
    'CV Matches': [3, 5, 2, 7, 4, 6, 5],
    'CVs Generated': [1, 0, 2, 1, 0, 1, 2],
    'Questions Practiced': [2, 3, 1, 4, 2, 3, 5]
})

fig_usage = go.Figure()

fig_usage.add_trace(go.Bar(
    x=usage_data['Date'],
    y=usage_data['CV Matches'],
    name='CV Matches',
    marker_color='#3b82f6',
    text=usage_data['CV Matches'],
    textposition='outside',
    textfont=dict(color='#ffffff')
))

fig_usage.add_trace(go.Bar(
    x=usage_data['Date'],
    y=usage_data['CVs Generated'],
    name='CVs Generated',
    marker_color='#10b981',
    text=usage_data['CVs Generated'],
    textposition='outside',
    textfont=dict(color='#ffffff')
))

fig_usage.add_trace(go.Bar(
    x=usage_data['Date'],
    y=usage_data['Questions Practiced'],
    name='Questions Practiced',
    marker_color='#8b5cf6',
    text=usage_data['Questions Practiced'],
    textposition='outside',
    textfont=dict(color='#ffffff')
))

fig_usage.update_layout(
    barmode='group',
    plot_bgcolor='rgba(0,0,0,0)',
    paper_bgcolor='rgba(0,0,0,0)',
    font=dict(color='#ffffff', size=14),
    xaxis=dict(
        showgrid=False,
        tickfont=dict(color='#ffffff', size=14)
    ),
    yaxis=dict(
        showgrid=True,
        gridcolor='rgba(255,255,255,0.1)',
        title='Count',
        titlefont=dict(color='#ffffff', size=16)
    ),
    legend=dict(
        orientation="h",
        yanchor="bottom",
        y=1.02,
        xanchor="right",
        x=1,
        font=dict(color='#ffffff', size=14)
    ),
    height=400
)

st.plotly_chart(fig_usage, use_container_width=True)

# ============================================================================
# EXPORT DATA
# ============================================================================
st.markdown("---")
st.markdown("### 📥 Export Analytics Report")

st.info("💡 Download your analytics data as a comprehensive report to track progress or share with mentors!")

export_col1, export_col2, export_col3 = st.columns(3)

with export_col1:
    # JSON Export
    export_data = {
        'generated_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'overview': {
            'total_matches': data['total_matches'],
            'avg_score': data['avg_score'],
            'cvs_generated': data['cvs_generated'],
            'interviews_practiced': data['interviews_practiced']
        },
        'match_history': data['match_history'],
        'skill_gaps': data['skill_gaps']
    }
    
    st.download_button(
        label="📄 JSON",
        data=json.dumps(export_data, indent=2),
        file_name=f"analytics_report_{datetime.now().strftime('%Y%m%d')}.json",
        mime="application/json",
        use_container_width=True
    )

with export_col2:
    # DOCX Export
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('CareerLens AI - Analytics Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph()
        
        # Overview Section
        doc.add_heading('Overview Metrics', level=1)
        
        overview_table = doc.add_table(rows=5, cols=2)
        overview_table.style = 'Light Grid Accent 1'
        
        overview_table.rows[0].cells[0].text = 'Metric'
        overview_table.rows[0].cells[1].text = 'Value'
        overview_table.rows[1].cells[0].text = 'Total CV Matches'
        overview_table.rows[1].cells[1].text = str(data['total_matches'])
        overview_table.rows[2].cells[0].text = 'Average Match Score'
        overview_table.rows[2].cells[1].text = f"{data['avg_score']}%"
        overview_table.rows[3].cells[0].text = 'CVs Generated'
        overview_table.rows[3].cells[1].text = str(data['cvs_generated'])
        overview_table.rows[4].cells[0].text = 'Interview Questions Practiced'
        overview_table.rows[4].cells[1].text = str(data['interviews_practiced'])
        
        doc.add_paragraph()
        
        # Match Score Trends
        doc.add_heading('Match Score Trends (Last 30 Days)', level=1)
        
        doc.add_paragraph(f"Latest Score: {match_df.iloc[-1]['score']}%")
        doc.add_paragraph(f"Highest Score: {match_df['score'].max()}%")
        doc.add_paragraph(f"Lowest Score: {match_df['score'].min()}%")
        doc.add_paragraph(f"Average Score: {match_df['score'].mean():.1f}%")
        
        doc.add_paragraph()
        
        # Skill Gaps
        doc.add_heading('Top Skill Gaps', level=1)
        
        doc.add_paragraph("These skills are most frequently missing from your CV matches:")
        doc.add_paragraph()
        
        for i, row in skill_gap_df.iterrows():
            doc.add_paragraph(f"{i+1}. {row['skill']} - Missing {row['frequency']} times", style='List Bullet')
        
        doc.add_paragraph()
        
        # Recommendations
        doc.add_heading('Recommendations', level=1)
        
        doc.add_paragraph("Based on your analytics:", style='List Bullet')
        doc.add_paragraph("Focus on learning Docker, Kubernetes, and AWS", style='List Bullet 2')
        doc.add_paragraph("Practice interview questions regularly", style='List Bullet 2')
        doc.add_paragraph("Aim for match scores above 75%", style='List Bullet 2')
        
        # Save to bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()
        
        st.download_button(
            label="📝 DOCX",
            data=docx_bytes,
            file_name=f"analytics_report_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
    except Exception as e:
        st.warning("⚠️ DOCX export unavailable")

with export_col3:
    # PDF Generation
    if st.button("🔄 Generate PDF", use_container_width=True, key="pdf_analytics_btn"):
        st.info("PDF conversion requires Microsoft Word. Download DOCX and convert manually, or use Google Docs → Download as PDF")

# Help section
with st.expander("ℹ️ Understanding Your Analytics"):
    st.markdown("""
    ### 📊 About This Dashboard
    
    **Current Status:** Demo Data  
    This dashboard currently shows simulated data for demonstration purposes.
    
    **In Production:**
    - Real-time tracking of all CV matches
    - Actual CV generation statistics
    - Live interview prep progress
    - Historical trend analysis
    
    **How Data Will Be Collected:**
    - Every CV-JD match automatically tracked
    - CV generations logged with metadata
    - Interview questions and evaluations recorded
    - Skill gaps identified from actual matches
    
    ---
    
    ### 🎯 Tips for Improvement
    
    1. **Maintain Consistency:** Use the platform 3-4 times per week
    2. **Focus on Gaps:** Prioritize learning high-frequency missing skills
    3. **Practice Regularly:** Evaluate 2-3 interview answers per week
    4. **Track Progress:** Check analytics weekly to measure improvement
    5. **Set Goals:** Aim for 80%+ match scores and evaluation scores
    """)