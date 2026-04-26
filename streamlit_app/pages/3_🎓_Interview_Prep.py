"""
Interview Prep Page - STAR Method Practice
"""
import streamlit as st
import sys
from pathlib import Path
import json
import os
from datetime import datetime
st.set_page_config(
    page_title="Interview Prep - CareerLens AI",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)
# Mobile responsiveness - Day 25
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
#from utils.mobile_styles import inject_mobile_styles

# Add project root to path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from src.guidance.interview_guidance import InterviewGuidance

# Page config

# Mobile responsiveness - Day 25
from mobile_styles import inject_mobile_styles
inject_mobile_styles()

# FIX: Force responsive layout on Cloudflare
st.markdown("""
<meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
<style>
    /* Force container width on desktop */
    .main .block-container {
        max-width: 100% !important;
        padding-left: 5rem !important;
        padding-right: 5rem !important;
    }
    
    /* Mobile override */
    @media (max-width: 768px) {
        .main .block-container {
            padding-left: 1rem !important;
            padding-right: 1rem !important;
        }
    }
    
    /* Fix Cloudflare font rendering */
    * {
        -webkit-font-smoothing: antialiased;
        -moz-osx-font-smoothing: grayscale;
        text-rendering: optimizeLegibility;
    }
    
    /* Prevent layout shift */
    .stApp {
        min-height: 100vh;
    }
    
    /* Fix column widths on desktop */
    [data-testid="column"] {
        width: auto !important;
        min-width: 0 !important;
    }
</style>
""", unsafe_allow_html=True)
# Dark theme CSS
st.markdown("""
<style>
    /* Force dark theme */
    .stApp {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%) !important;
    }
    
    /* Hide branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    /* All text light colored */
    * {
        color: #e2e8f0 !important;
    }
    
    /* Sidebar */
    [data-testid="stSidebar"] {
        background: rgba(15, 23, 42, 0.95) !important;
    }
    
    /* Text inputs */
    .stTextInput input, .stTextArea textarea {
        background: rgba(30, 41, 59, 0.8) !important;
        color: #e2e8f0 !important;
        border: 1px solid rgba(96, 165, 250, 0.3) !important;
        border-radius: 8px;
    }
    
    /* Buttons */
    .stButton > button {
        background: linear-gradient(135deg, #3b82f6, #8b5cf6) !important;
        color: white !important;
        border: none;
        padding: 0.875rem 2rem;
        font-size: 1rem;
        border-radius: 12px;
        font-weight: 600;
        box-shadow: 0 4px 15px rgba(59, 130, 246, 0.4);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(59, 130, 246, 0.6);
    }
    
    /* Question cards */
    .question-card {
        background: rgba(30, 41, 59, 0.7);
        border-left: 4px solid #3b82f6;
        border-radius: 8px;
        padding: 1.5rem;
        margin-bottom: 1rem;
        transition: all 0.3s ease;
    }
    
    .question-card:hover {
        border-left-color: #60a5fa;
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.2);
    }
    
    .question-difficulty-easy {
        border-left-color: #10b981;
    }
    
    .question-difficulty-medium {
        border-left-color: #f59e0b;
    }
    
    .question-difficulty-hard {
        border-left-color: #ef4444;
    }
    
    /* STAR framework box */
    .star-box {
        background: linear-gradient(135deg, rgba(59, 130, 246, 0.1), rgba(139, 92, 246, 0.1));
        border: 1px solid rgba(96, 165, 250, 0.3);
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1rem 0;
    }
    
    .star-component {
        background: rgba(30, 41, 59, 0.5);
        border-left: 3px solid #60a5fa;
        border-radius: 6px;
        padding: 1rem;
        margin: 0.75rem 0;
    }
    
    /* Score display */
    .score-display {
        background: linear-gradient(135deg, rgba(16, 185, 129, 0.1), rgba(52, 211, 153, 0.1));
        border: 2px solid rgba(16, 185, 129, 0.3);
        border-radius: 12px;
        padding: 2rem;
        text-align: center;
        margin: 1rem 0;
    }
    
    .score-value {
        font-size: 3rem;
        font-weight: 800;
        background: linear-gradient(135deg, #10b981, #34d399);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    
    /* Feedback boxes */
    .feedback-good {
        background: rgba(16, 185, 129, 0.1);
        border-left: 3px solid #10b981;
        padding: 0.75rem;
        margin: 0.5rem 0;
        border-radius: 4px;
    }
    
    .feedback-warning {
        background: rgba(245, 158, 11, 0.1);
        border-left: 3px solid #f59e0b;
        padding: 0.75rem;
        margin: 0.5rem 0;
        border-radius: 4px;
    }
    
    /* Success/Info messages */
    .stSuccess {
        background: rgba(16, 185, 129, 0.1) !important;
        border: 1px solid rgba(16, 185, 129, 0.3) !important;
        color: #10b981 !important;
    }
    
    .stInfo {
        background: rgba(59, 130, 246, 0.1) !important;
        border: 1px solid rgba(59, 130, 246, 0.3) !important;
        color: #3b82f6 !important;
    }
    
    .stWarning {
        background: rgba(245, 158, 11, 0.1) !important;
        border: 1px solid rgba(245, 158, 11, 0.3) !important;
        color: #f59e0b !important;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'interview_system' not in st.session_state:
    st.session_state.interview_system = InterviewGuidance()
if 'selected_question' not in st.session_state:
    st.session_state.selected_question = None
if 'star_answer' not in st.session_state:
    st.session_state.star_answer = None
if 'evaluation_result' not in st.session_state:
    st.session_state.evaluation_result = None

# Title
st.markdown("""
<h1 style="background: linear-gradient(90deg, #60a5fa, #a78bfa); -webkit-background-clip: text; 
           -webkit-text-fill-color: transparent; font-size: 3rem; font-weight: 800; margin-bottom: 0.5rem;">
    🎓 Interview Preparation
</h1>
<p style="color: #94a3b8 !important; font-size: 1.25rem; margin-bottom: 2rem;">
    Master the STAR method • Practice with AI feedback • Build confidence
</p>
""", unsafe_allow_html=True)

# Two modes: Practice Questions or STAR Builder
mode = st.radio(
    "Choose your practice mode:",
    ["🎯 Get Interview Questions", "📊 Evaluate My Answer"],
    horizontal=True
)

st.markdown("---")

# ============================================================================
# MODE 1: GET INTERVIEW QUESTIONS
# ============================================================================
if "Get Interview Questions" in mode:
    st.markdown("### 🎯 Get Personalized Interview Questions")
    
    st.info("💡 Enter your skills and we'll generate relevant interview questions across different categories!")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        skills_input = st.text_input(
            "Enter your skills (comma-separated)",
            placeholder="python, fastapi, docker, kubernetes, react, aws",
            help="List the skills you want to be interviewed on"
        )
    
    with col2:
        num_questions = st.slider("Number of questions", min_value=5, max_value=20, value=10, step=1)
    
    if st.button("🎯 Generate Questions", type="primary", use_container_width=True):
        if not skills_input:
            st.error("❌ Please enter at least one skill")
            st.stop()
        
        skills = [s.strip() for s in skills_input.split(',') if s.strip()]
        
        with st.spinner("🧠 Generating personalized interview questions..."):
            try:
                interview_system = st.session_state.interview_system
                questions = interview_system.generate_questions(
                    skills=skills,
                    num_questions=num_questions
                )
                
                #st.success(f"✅ Generated {questions['total_questions']} questions across {len(questions['by_category'])} categories!")
                # Get total questions safely
                total_questions = len(questions.get('questions', []))
                st.success(f"✅ Generated {total_questions} interview questions!")
                # Display questions by category
                st.markdown("### 📝 Your Interview Questions")
                
                for category, cat_questions in questions['by_category'].items():
                    if cat_questions:
                        st.markdown(f"#### {category.replace('_', ' ').title()} ({len(cat_questions)} questions)")
                        
                        for i, q in enumerate(cat_questions, 1):
                            difficulty = q['difficulty'].lower()
                            difficulty_class = f"question-difficulty-{difficulty}"
                            
                            # Difficulty emoji
                            diff_emoji = "🟢" if difficulty == "easy" else ("🟡" if difficulty == "medium" else "🔴")
                            
                            st.markdown(f"""
                            <div class="question-card {difficulty_class}">
                                <strong style="font-size: 1.1rem;">{i}. {q['question']}</strong>
                                <div style="margin-top: 0.5rem; color: #94a3b8; font-size: 0.9rem;">
                                    {diff_emoji} Difficulty: {q['difficulty']} • Category: {q['category']}
                                </div>
                            </div>
                            """, unsafe_allow_html=True)
                            
                            # Show hint if available
                            if 'hint' in q and q['hint']:
                                with st.expander(f"💡 Hint for Question {i}"):
                                    st.write(q['hint'])
                
                # Download options
                st.markdown("---")
                st.markdown("### 📥 Export Questions")
                
                export_col1, export_col2, export_col3 = st.columns(3)
                
                with export_col1:
                    # JSON Download
                    st.download_button(
                        label="📄 JSON",
                        data=json.dumps(questions, indent=2),
                        file_name="interview_questions.json",
                        mime="application/json",
                        use_container_width=True
                    )
                
                with export_col2:
                    # DOCX Download
                    try:
                        from docx import Document
                        from docx.shared import Pt
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        import io
                        
                        # Create document
                        doc = Document()
                        
                        # Title
                        title = doc.add_heading('Interview Questions', 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add metadata
                        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
                        doc.add_paragraph(f"Total Questions: {questions['total_questions']}")
                        doc.add_paragraph(f"Skills: {', '.join(skills)}")
                        doc.add_paragraph()
                        
                        # Add questions by category
                        for category, cat_questions in questions['by_category'].items():
                            if cat_questions:
                                # Category heading
                                doc.add_heading(f"{category} ({len(cat_questions)} questions)", level=1)
                                
                                for i, q in enumerate(cat_questions, 1):
                                    # Question number and text
                                    p = doc.add_paragraph()
                                    p.add_run(f"Q{i}. ").bold = True
                                    p.add_run(q['question'])
                                    
                                    # Difficulty
                                    diff_p = doc.add_paragraph()
                                    diff_run = diff_p.add_run(f"Difficulty: {q['difficulty']}")
                                    diff_run.font.size = Pt(10)
                                    
                                    # Hint if available
                                    if q.get('hint'):
                                        hint_p = doc.add_paragraph()
                                        hint_run = hint_p.add_run(f"Hint: {q['hint']}")
                                        hint_run.font.size = Pt(10)
                                        hint_run.font.italic = True
                                    
                                    doc.add_paragraph()  # Spacing
                        
                        # Add preparation tips
                        doc.add_page_break()
                        doc.add_heading('Preparation Tips', level=1)
                        for tip in questions['preparation_tips']:
                            doc.add_paragraph(tip, style='List Bullet')
                        
                        # Save to bytes
                        docx_buffer = io.BytesIO()
                        doc.save(docx_buffer)
                        docx_bytes = docx_buffer.getvalue()
                        
                        st.download_button(
                            label="📝 DOCX",
                            data=docx_bytes,
                            file_name=f"interview_questions_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        
                    except Exception as e:
                        st.warning("⚠️ DOCX export unavailable. Download JSON instead.")
                
                with export_col3:
                    # PDF Generation Button
                    if st.button("🔄 Generate PDF", use_container_width=True, key="pdf_questions_btn"):
                        try:
                            from docx import Document
                            from docx.shared import Pt
                            from docx.enum.text import WD_ALIGN_PARAGRAPH
                            import io
                            import tempfile
                            
                            # Create DOCX first
                            doc = Document()
                            
                            # Title
                            title = doc.add_heading('Interview Questions', 0)
                            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Add metadata
                            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
                            doc.add_paragraph(f"Total Questions: {questions['total_questions']}")
                            doc.add_paragraph(f"Skills: {', '.join(skills)}")
                            doc.add_paragraph()
                            
                            # Add questions by category
                            for category, cat_questions in questions['by_category'].items():
                                if cat_questions:
                                    doc.add_heading(f"{category} ({len(cat_questions)} questions)", level=1)
                                    
                                    for i, q in enumerate(cat_questions, 1):
                                        p = doc.add_paragraph()
                                        p.add_run(f"Q{i}. ").bold = True
                                        p.add_run(q['question'])
                                        
                                        diff_p = doc.add_paragraph()
                                        diff_run = diff_p.add_run(f"Difficulty: {q['difficulty']}")
                                        diff_run.font.size = Pt(10)
                                        
                                        if q.get('hint'):
                                            hint_p = doc.add_paragraph()
                                            hint_run = hint_p.add_run(f"Hint: {q['hint']}")
                                            hint_run.font.size = Pt(10)
                                            hint_run.font.italic = True
                                        
                                        doc.add_paragraph()
                            
                            # Add preparation tips
                            doc.add_page_break()
                            doc.add_heading('Preparation Tips', level=1)
                            for tip in questions['preparation_tips']:
                                doc.add_paragraph(tip, style='List Bullet')
                            
                            # Save DOCX temporarily
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                                doc.save(tmp_docx.name)
                                tmp_docx_path = tmp_docx.name
                            
                            # Try to convert to PDF
                            try:
                                from docx2pdf import convert
                                
                                pdf_path = tmp_docx_path.replace('.docx', '.pdf')
                                convert(tmp_docx_path, pdf_path)
                                
                                with open(pdf_path, 'rb') as pdf_file:
                                    pdf_bytes = pdf_file.read()
                                
                                # Clean up
                                os.unlink(tmp_docx_path)
                                os.unlink(pdf_path)
                                
                                # Store in session state
                                st.session_state.questions_pdf = pdf_bytes
                                st.success("✅ PDF generated!")
                                st.rerun()
                                
                            except ImportError:
                                st.warning("""
                                ⚠️ PDF conversion not available.
                                
                                **Options:**
                                1. Download DOCX and convert using Word
                                2. Download DOCX and upload to Google Docs → Download as PDF
                                """)
                                os.unlink(tmp_docx_path)
                                
                        except Exception as e:
                            st.error(f"❌ PDF generation failed: {str(e)}")
                
                # Show PDF download if available
                if 'questions_pdf' in st.session_state and st.session_state.questions_pdf:
                    st.markdown("---")
                    st.download_button(
                        label="📕 Download PDF",
                        data=st.session_state.questions_pdf,
                        file_name=f"interview_questions_{datetime.now().strftime('%Y%m%d')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                
            except Exception as e:
                st.error(f"❌ Failed to generate questions: {str(e)}")
                st.exception(e)


# ============================================================================
# MODE 3: EVALUATE ANSWER
# ============================================================================
elif "Evaluate My Answer" in mode:
    st.markdown("### 📊 Evaluate Your Interview Answer")
    
    st.info("💡 Paste your interview answer below and get AI-powered feedback based on the STAR method!")
    
    # Question input
    eval_question = st.text_area(
        "Interview Question:",
        placeholder="Tell me about a time when...",
        height=80
    )
    
    # Answer input
    eval_answer = st.text_area(
        "Your Answer:",
        placeholder="In my previous role as...",
        height=300,
        help="Write your complete interview answer (200-500 words recommended)"
    )
    
    col1, col2 = st.columns(2)
    
    with col1:
        word_count = len(eval_answer.split()) if eval_answer else 0
        st.metric("Word Count", word_count, 
                 "✅ Good length" if 200 <= word_count <= 500 else "⚠️ Too short" if word_count < 200 else "⚠️ Too long")
    
    with col2:
        if st.button("📊 Evaluate Answer", type="primary", use_container_width=True):
            if not eval_question or not eval_answer:
                st.error("❌ Please provide both question and answer")
                st.stop()
            
            with st.spinner("🧠 Evaluating your answer..."):
                try:
                    interview_system = st.session_state.interview_system
                    evaluation = interview_system.evaluate_answer(eval_question, eval_answer)
                    
                    st.session_state.evaluation_result = evaluation
                    
                    st.success("✅ Evaluation complete!")
                    
                except Exception as e:
                    st.error(f"❌ Evaluation failed: {str(e)}")
    
    # Display evaluation results
    if st.session_state.evaluation_result:
        st.markdown("---")
        st.markdown("## 🎯 Evaluation Results")
        
        result = st.session_state.evaluation_result
        
        # Overall score
        score = result.get('overall_score', result.get('score', 0))
        rating = result.get('rating', 'Evaluated')
        
        # Color based on rating
        if rating == "Excellent":
            color = "#10b981"
            emoji = "🌟"
        elif rating == "Good":
            color = "#3b82f6"
            emoji = "✅"
        elif rating == "Needs Improvement":
            color = "#f59e0b"
            emoji = "⚠️"
        else:
            color = "#ef4444"
            emoji = "❌"
        
        st.markdown(f"""
        <div class="score-display" style="border-color: {color};">
            <div class="score-value" style="background: linear-gradient(135deg, {color}, {color}); 
                 -webkit-background-clip: text; -webkit-text-fill-color: transparent;">
                {emoji} {score}/100
            </div>
            <div style="font-size: 1.5rem; color: {color}; margin-top: 0.5rem;">
                {rating}
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Breakdown
        st.markdown("### 📊 Score Breakdown")
        
        breakdown_col1, breakdown_col2 = st.columns(2)
        
        with breakdown_col1:
            st.markdown("#### STAR Components")
            
            for component, details in result['breakdown'].items():
                if component != 'length':
                    score_val = details['score']
                    max_score = details['max_score']
                    percentage = (score_val / max_score * 100) if max_score > 0 else 0
                    
                    st.metric(
                        component.title(),
                        f"{score_val}/{max_score}",
                        f"{percentage:.0f}%"
                    )
                    st.progress(percentage / 100)
        
        with breakdown_col2:
            st.markdown("#### Length Check")
            
            length_details = result['breakdown']['length']
            st.metric(
                "Word Count",
                length_details['word_count'],
                length_details['status']
            )
            st.progress(min(length_details['score'] / 20, 1.0))
        
        # Feedback
        st.markdown("### 💡 Detailed Feedback")
        
        for feedback_item in result['feedback']:
            if "✅" in feedback_item or "Strong" in feedback_item:
                st.markdown(f"""
                <div class="feedback-good">
                    {feedback_item}
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="feedback-warning">
                    {feedback_item}
                </div>
                """, unsafe_allow_html=True)
        
        # Suggestions
        if result.get('suggestions'):
            st.markdown("### 🎯 Improvement Suggestions")
            
            for i, suggestion in enumerate(result['suggestions'], 1):
                st.warning(f"{i}. {suggestion}")

# Help section
with st.expander("ℹ️ Interview Preparation Tips"):
    st.markdown("""
    ### 🎯 STAR Method Explained
    
    The STAR method helps you structure behavioral interview answers:
    
    **S - Situation (20-25% of answer)**
    - Set the scene and give context
    - When and where did this happen?
    - What was the challenge or opportunity?
    
    **T - Task (20-25% of answer)**
    - What was your specific role/responsibility?
    - What goal were you working toward?
    - What were the stakes?
    
    **A - Action (30-35% of answer)**
    - What specific steps did YOU take?
    - What decisions did you make?
    - How did you solve the problem?
    - Focus on "I" not "we"
    
    **R - Result (20-25% of answer)**
    - What was the outcome?
    - Include quantifiable metrics
    - What did you learn?
    - How did it benefit the company/team?
    
    ---
    
    ### 💡 Best Practices
    
    **Before the Interview:**
    - ✅ Prepare 5-7 STAR stories covering different competencies
    - ✅ Practice out loud (not just in your head)
    - ✅ Time yourself (aim for 2-3 minutes per answer)
    - ✅ Quantify results whenever possible
    
    **During the Interview:**
    - ✅ Take a moment to think before answering
    - ✅ Speak clearly and confidently
    - ✅ Make eye contact (or camera contact for video)
    - ✅ Use specific examples, not generalities
    
    **Common Mistakes to Avoid:**
    - ❌ Being too vague or generic
    - ❌ Talking about "we" instead of "I"
    - ❌ Forgetting to mention results
    - ❌ Making your answer too long (>3 minutes)
    - ❌ Not preparing enough examples in advance
    
    ---
    
    ### 📚 Question Categories
    
    **Behavioral Questions**
    - Leadership, teamwork, problem-solving
    - Use STAR method
    - Example: "Tell me about a time when..."
    
    **Technical Questions**
    - Skill-specific knowledge
    - Explain concepts clearly
    - Example: "What is the difference between..."
    
    **Coding Questions**
    - Live problem-solving
    - Think out loud
    - Example: "Write a function to..."
    
    **System Design Questions**
    - Architecture and scalability
    - Discuss trade-offs
    - Example: "How would you design..."
    """)