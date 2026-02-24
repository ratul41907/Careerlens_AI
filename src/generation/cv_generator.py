"""
CV Generator - Creates ATS-optimized CVs in DOCX format
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Optional
from datetime import datetime
from loguru import logger


class CVGenerator:
    """
    Generate ATS-optimized CVs in DOCX format
    """
    
    def __init__(self):
        """Initialize CV generator"""
        logger.info("CVGenerator initialized")
        
        # ATS-friendly settings
        self.font_name = "Calibri"  # ATS-safe font
        self.font_size = 11
        self.heading_size = 14
        self.name_size = 18
        
    def generate_cv(self,
                   personal_info: Dict,
                   experience: List[Dict],
                   education: List[Dict],
                   skills: List[str],
                   projects: Optional[List[Dict]] = None,
                   certifications: Optional[List[str]] = None,
                   target_jd: Optional[Dict] = None) -> Document:
        """
        Generate ATS-optimized CV
        
        Args:
            personal_info: {'name': str, 'email': str, 'phone': str, 'location': str, 'summary': str}
            experience: List of {'title': str, 'company': str, 'duration': str, 'bullets': List[str]}
            education: List of {'degree': str, 'institution': str, 'year': str, 'gpa': str (optional)}
            skills: List of skill strings
            projects: Optional list of {'name': str, 'description': str, 'technologies': str}
            certifications: Optional list of certification strings
            target_jd: Optional JD data for skill highlighting
            
        Returns:
            python-docx Document object
        """
        logger.info("Generating ATS-optimized CV...")
        
        # Create document
        doc = Document()
        
        # Set margins (1 inch all sides - ATS standard)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add sections
        self._add_header(doc, personal_info)
        
        if personal_info.get('summary'):
            self._add_summary(doc, personal_info['summary'])
        
        if skills:
            self._add_skills(doc, skills, target_jd)
        
        if experience:
            self._add_experience(doc, experience)
        
        if education:
            self._add_education(doc, education)
        
        if projects:
            self._add_projects(doc, projects)
        
        if certifications:
            self._add_certifications(doc, certifications)
        
        logger.info("CV generated successfully")
        return doc
    
    def _add_header(self, doc: Document, personal_info: Dict):
        """Add name and contact information"""
        # Name (centered, large)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal_info['name'].upper())
        name_run.font.size = Pt(self.name_size)
        name_run.font.name = self.font_name
        name_run.bold = True
        
        # Contact info (centered, one line)
        contact_parts = []
        if personal_info.get('email'):
            contact_parts.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_parts.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_parts.append(personal_info['location'])
        if personal_info.get('linkedin'):
            contact_parts.append(personal_info['linkedin'])
        
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(self.font_size)
        contact_run.font.name = self.font_name
        
        # Spacing
        doc.add_paragraph()
    
    def _add_summary(self, doc: Document, summary: str):
        """Add professional summary"""
        # Section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PROFESSIONAL SUMMARY")
        heading_run.font.size = Pt(self.heading_size)
        heading_run.font.name = self.font_name
        heading_run.bold = True
        
        # Underline (ATS-safe separator)
        separator = doc.add_paragraph()
        sep_run = separator.add_run("_" * 80)
        sep_run.font.size = Pt(8)
        
        # Summary text
        summary_para = doc.add_paragraph(summary)
        summary_para.paragraph_format.space_after = Pt(12)
        self._format_paragraph(summary_para)
        
        doc.add_paragraph()
    
    def _add_skills(self, doc: Document, skills: List[str], target_jd: Optional[Dict] = None):
        """Add skills section with optional JD-based highlighting"""
        # Section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("TECHNICAL SKILLS")
        heading_run.font.size = Pt(self.heading_size)
        heading_run.font.name = self.font_name
        heading_run.bold = True
        
        # Underline
        separator = doc.add_paragraph()
        sep_run = separator.add_run("_" * 80)
        sep_run.font.size = Pt(8)
        
        # Organize skills by category (if possible)
        # For now, simple comma-separated list
        skills_para = doc.add_paragraph()
        
        # If target JD provided, highlight matching skills
        if target_jd and target_jd.get('required_skills'):
            required_skills_lower = [s.lower() for s in target_jd['required_skills']]
            
            for i, skill in enumerate(skills):
                skill_run = skills_para.add_run(skill)
                skill_run.font.name = self.font_name
                skill_run.font.size = Pt(self.font_size)
                
                # Bold if matches JD requirements
                if skill.lower() in required_skills_lower:
                    skill_run.bold = True
                
                # Add separator
                if i < len(skills) - 1:
                    skills_para.add_run(" • ")
        else:
            # No highlighting
            skills_text = " • ".join(skills)
            skills_run = skills_para.add_run(skills_text)
            self._format_run(skills_run)
        
        skills_para.paragraph_format.space_after = Pt(12)
        doc.add_paragraph()
    
    def _add_experience(self, doc: Document, experience: List[Dict]):
        """Add work experience section"""
        # Section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PROFESSIONAL EXPERIENCE")
        heading_run.font.size = Pt(self.heading_size)
        heading_run.font.name = self.font_name
        heading_run.bold = True
        
        # Underline
        separator = doc.add_paragraph()
        sep_run = separator.add_run("_" * 80)
        sep_run.font.size = Pt(8)
        
        for exp in experience:
            # Job title and company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{exp['title']} | {exp['company']}")
            title_run.font.name = self.font_name
            title_run.font.size = Pt(self.font_size)
            title_run.bold = True
            
            # Duration
            duration_para = doc.add_paragraph(exp['duration'])
            duration_para.paragraph_format.space_after = Pt(6)
            self._format_paragraph(duration_para)
            duration_para.runs[0].italic = True
            
            # Bullets
            for bullet in exp['bullets']:
                bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                self._format_paragraph(bullet_para)
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                bullet_para.paragraph_format.space_after = Pt(3)
            
            # Spacing between jobs
            doc.add_paragraph()
    
    def _add_education(self, doc: Document, education: List[Dict]):
        """Add education section"""
        # Section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("EDUCATION")
        heading_run.font.size = Pt(self.heading_size)
        heading_run.font.name = self.font_name
        heading_run.bold = True
        
        # Underline
        separator = doc.add_paragraph()
        sep_run = separator.add_run("_" * 80)
        sep_run.font.size = Pt(8)
        
        for edu in education:
            # Degree and institution
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(f"{edu['degree']} | {edu['institution']}")
            degree_run.font.name = self.font_name
            degree_run.font.size = Pt(self.font_size)
            degree_run.bold = True
            
            # Year and GPA
            details = [edu['year']]
            if edu.get('gpa'):
                details.append(f"GPA: {edu['gpa']}")
            
            details_para = doc.add_paragraph(" | ".join(details))
            details_para.paragraph_format.space_after = Pt(6)
            self._format_paragraph(details_para)
            details_para.runs[0].italic = True
        
        doc.add_paragraph()
    
    def _add_projects(self, doc: Document, projects: List[Dict]):
        """Add projects section"""
        # Section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("PROJECTS")
        heading_run.font.size = Pt(self.heading_size)
        heading_run.font.name = self.font_name
        heading_run.bold = True
        
        # Underline
        separator = doc.add_paragraph()
        sep_run = separator.add_run("_" * 80)
        sep_run.font.size = Pt(8)
        
        for project in projects:
            # Project name
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(project['name'])
            name_run.font.name = self.font_name
            name_run.font.size = Pt(self.font_size)
            name_run.bold = True
            
            # Description
            desc_para = doc.add_paragraph(project['description'])
            desc_para.paragraph_format.left_indent = Inches(0.25)
            self._format_paragraph(desc_para)
            
            # Technologies
            if project.get('technologies'):
                tech_para = doc.add_paragraph(f"Technologies: {project['technologies']}")
                tech_para.paragraph_format.left_indent = Inches(0.25)
                tech_para.paragraph_format.space_after = Pt(6)
                self._format_paragraph(tech_para)
                tech_para.runs[0].italic = True
            
            doc.add_paragraph()
    
    def _add_certifications(self, doc: Document, certifications: List[str]):
        """Add certifications section"""
        # Section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("CERTIFICATIONS")
        heading_run.font.size = Pt(self.heading_size)
        heading_run.font.name = self.font_name
        heading_run.bold = True
        
        # Underline
        separator = doc.add_paragraph()
        sep_run = separator.add_run("_" * 80)
        sep_run.font.size = Pt(8)
        
        # List certifications
        for cert in certifications:
            cert_para = doc.add_paragraph(cert, style='List Bullet')
            self._format_paragraph(cert_para)
            cert_para.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_paragraph()
    
    def _format_paragraph(self, paragraph):
        """Apply standard formatting to paragraph"""
        for run in paragraph.runs:
            self._format_run(run)
    
    def _format_run(self, run):
        """Apply standard formatting to text run"""
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
    
    def save_cv(self, doc: Document, filepath: str):
        """Save CV to file"""
        doc.save(filepath)
        logger.info(f"CV saved to: {filepath}")


# Convenience function
def generate_cv(personal_info: Dict,
               experience: List[Dict],
               education: List[Dict],
               skills: List[str],
               **kwargs) -> Document:
    """Quick CV generation function"""
    generator = CVGenerator()
    return generator.generate_cv(
        personal_info, experience, education, skills, **kwargs
    )